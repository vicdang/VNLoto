import random
import json
import os
import time
import argparse
from concurrent.futures import ProcessPoolExecutor, as_completed

from docx import Document
from docx.shared import Pt, Cm, Mm, Inches, RGBColor, Emu
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# Global configuration - will be set from config.json
NUM_COLUMNS = 9
COLUMN_RANGES = []

def init_column_ranges(num_cols):
    """Generate column ranges for Lô Tô based on number of columns.
    9 columns: 1-10, 10-20, ..., 80-90
    8 columns: 1-10-1, 10-20, ..., 70-81 (adjusted for 80 numbers)
    """
    global NUM_COLUMNS, COLUMN_RANGES
    NUM_COLUMNS = num_cols
    COLUMN_RANGES = []
    numbers_per_col = 90 // num_cols
    for i in range(num_cols):
        start = i * numbers_per_col + 1
        end = (i + 1) * numbers_per_col + 1
        COLUMN_RANGES.append(list(range(start, end)))

# Default initialization for 9 columns
init_column_ranges(9)

CARDS_PER_TABLE = 3
NUM_WORKERS = os.cpu_count() or 4


def generate_card(used_numbers=None):
    if used_numbers is None:
        used_numbers = set()

    while True:
        # Step 1: Decide how many numbers per column (1, 2, or 3), total must be 15
        while True:
            col_counts = [random.randint(1, 3) for _ in range(NUM_COLUMNS)]
            if sum(col_counts) == 15:
                break

        # Step 2: Pick numbers for each column, avoiding used_numbers as much as possible
        col_numbers = []
        card_nums = set()
        valid = True
        for c in range(NUM_COLUMNS):
            available = [n for n in COLUMN_RANGES[c] if n not in used_numbers]
            if len(available) < col_counts[c]:
                available = [n for n in COLUMN_RANGES[c] if n not in card_nums]
            if len(available) < col_counts[c]:
                valid = False
                break
            chosen = sorted(random.sample(available, col_counts[c]))
            col_numbers.append(chosen)
            card_nums.update(chosen)

        if not valid:
            continue

        # Step 3: Distribute numbers into 3 rows, each row must have exactly 5 numbers
        # Build a placement: for each column, decide which rows get numbers
        grid = [[None] * NUM_COLUMNS for _ in range(3)]

        # For each column, pick which rows get the numbers
        row_fills = [0, 0, 0]  # track how many numbers each row has
        col_row_assignments = []

        for c in range(NUM_COLUMNS):
            count = col_counts[c]
            if count == 3:
                rows = [0, 1, 2]
            elif count == 2:
                rows = sorted(random.sample([0, 1, 2], 2))
            else:
                rows = [random.choice([0, 1, 2])]
            col_row_assignments.append(rows)
            for r in rows:
                row_fills[r] += 1

        # Check if each row has exactly 5 — if not, retry
        if row_fills != [5, 5, 5]:
            # Try to fix by swapping assignments
            # Use a backtracking approach with limited attempts
            fixed = False
            for attempt in range(200):
                # Reset
                test_assignments = []
                test_fills = [0, 0, 0]
                ok = True
                
                # Shuffle column order for variety
                order = list(range(NUM_COLUMNS))
                random.shuffle(order)
                temp_assignments = [None] * NUM_COLUMNS
                
                for c in order:
                    count = col_counts[c]
                    if count == 3:
                        rows = [0, 1, 2]
                    elif count == 2:
                        # Pick 2 rows that have the most remaining capacity
                        remaining = [(5 - test_fills[r], r) for r in range(3)]
                        remaining.sort(reverse=True)
                        candidates = [r for cap, r in remaining if cap > 0]
                        if len(candidates) < 2:
                            ok = False
                            break
                        # Weighted random: prefer rows with more capacity
                        rows = sorted(random.sample(candidates, 2))
                    else:
                        remaining = [(5 - test_fills[r], r) for r in range(3)]
                        remaining.sort(reverse=True)
                        candidates = [r for cap, r in remaining if cap > 0]
                        if len(candidates) < 1:
                            ok = False
                            break
                        rows = [random.choice(candidates)]
                    
                    temp_assignments[c] = rows
                    for r in rows:
                        test_fills[r] += 1

                if ok and test_fills == [5, 5, 5]:
                    col_row_assignments = temp_assignments
                    fixed = True
                    break

            if not fixed:
                continue

        # Place numbers in grid
        for c in range(NUM_COLUMNS):
            nums = col_numbers[c]
            rows = col_row_assignments[c]
            for i, r in enumerate(rows):
                grid[r][c] = nums[i]

        # Validate
        if not validate_card(grid):
            continue

        used_numbers.update(card_nums)
        return grid


def validate_card(grid):
    # Rule: each row has exactly 5 numbers
    for r in range(3):
        count = sum(1 for v in grid[r] if v is not None)
        if count != 5:
            return False

    # Rule: total 15 numbers
    all_nums = [v for row in grid for v in row if v is not None]
    if len(all_nums) != 15:
        return False

    # Rule: no duplicates
    if len(set(all_nums)) != 15:
        return False

    # Rule: numbers in correct column ranges
    for r in range(3):
        for c in range(NUM_COLUMNS):
            v = grid[r][c]
            if v is not None:
                if v not in COLUMN_RANGES[c]:
                    return False

    # Rule: columns sorted ascending top to bottom
    for c in range(NUM_COLUMNS):
        col_vals = [grid[r][c] for r in range(3) if grid[r][c] is not None]
        if col_vals != sorted(col_vals):
            return False

    return True


def generate_table(_=None):
    """Generate one table (3 cards, 9 rows). Returns the 9-row grid."""
    used = set()
    rows = []
    for _ in range(CARDS_PER_TABLE):
        grid = generate_card(used)
        rows.extend(grid)
    return rows


def table_fingerprint(table):
    """Create a hashable fingerprint to detect duplicate tables."""
    return tuple(
        v for row in table for v in row
    )


def worker_batch(batch_size):
    """Worker function: generate a batch of unique tables."""
    tables = []
    seen = set()
    for _ in range(batch_size):
        while True:
            t = generate_table()
            fp = table_fingerprint(t)
            if fp not in seen:
                seen.add(fp)
                tables.append(t)
                break
    return tables


def print_combined_visual(table, table_id, file=None):
    border = "+" + "----+" * NUM_COLUMNS
    thick  = "+" + "====+" * NUM_COLUMNS
    out = file or __import__("sys").stdout
    out.write(f"--- Table {table_id:03d} ---\n")
    out.write(thick + "\n")
    for idx, row in enumerate(table):
        cells = []
        for v in row:
            if v is None:
                cells.append("  . ")
            else:
                cells.append(f" {v:2d} ")
        out.write("|" + "|".join(cells) + "|\n")
        # thick border between cards (every 3 rows), thin otherwise
        if (idx + 1) % 3 == 0:
            out.write(thick + "\n")
        else:
            out.write(border + "\n")


def validate_table(table):
    """Validate all 3 cards in a 9-row table."""
    for card_idx in range(CARDS_PER_TABLE):
        card = table[card_idx * 3 : card_idx * 3 + 3]
        if not validate_card(card):
            return False
    return True


def save_docx(all_tables, path, config, tables_per_page=8):
    """Save tables to DOCX with 2 tables per row, laid out in a grid."""
    doc = Document()
    title_text = config.get("title", "LÔ TÔ")
    font_name = config.get("font", "Calibri")
    font_size = Pt(config.get("font_size", 16))
    col_width = Inches(config.get("column_width_inches", 0.3))
    row_ht = Inches(config.get("row_height_inches", 0.3))
    header_font_size = Pt(config.get("header_font_size", 18))
    header_caps = config.get("header_text_caps", True)
    header_bg = config.get("header_bg_color", "808080")
    special_count = max(0, min(10, config.get("special_cells", 0)))
    special_char = config.get("special_cell_char", "★")
    special_replace = config.get("special_cell_replace", False)
    num_rounds = max(1, config.get("rounds", 1))
    total_tables_cfg = config.get("total_tables", len(all_tables))
    tables_per_round = max(1, total_tables_cfg // num_rounds)
    footer_messages = config.get("footer_messages", None)
    if not footer_messages:
        # Fallback: single string from old config or default
        fallback = config.get("footer_text", "Have a great trip, DC34 love you - PUB Team")
        footer_messages = [fallback]
    footer_font_size = Pt(config.get("footer_font_size", 10))
    border_cfg = config.get("table_border", {})
    border_style = border_cfg.get("style", "single")
    border_size = str(border_cfg.get("size", 12))
    border_color = border_cfg.get("color", "000000")

    COLS_PER_ROW = 4
    rows_per_page = (tables_per_page + COLS_PER_ROW - 1) // COLS_PER_ROW

    # --- Page setup: Landscape A4 ---
    for section in doc.sections:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Cm(29.7)
        section.page_height = Cm(21)
        section.top_margin = Cm(1.0)
        section.bottom_margin = Cm(1.0)
        section.left_margin = Cm(1.2)
        section.right_margin = Cm(1.2)

    # --- Border XML templates ---
    NO_BORDER_XML = (
        '<w:tcBorders %s>'
        '<w:top w:val="none" w:sz="0" w:space="0"/>'
        '<w:left w:val="none" w:sz="0" w:space="0"/>'
        '<w:right w:val="none" w:sz="0" w:space="0"/>'
        '<w:bottom w:val="none" w:sz="0" w:space="0"/>'
        '</w:tcBorders>' % nsdecls('w')
    )
    THIN_BORDER_XML = (
        '<w:tcBorders %s>'
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '</w:tcBorders>' % nsdecls('w')
    )
    CARD_BOTTOM_XML = (
        '<w:tcBorders %s>'
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:bottom w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
        '</w:tcBorders>' % nsdecls('w')
    )
    TITLE_BORDER_XML = (
        '<w:tcBorders %s>'
        '<w:top w:val="%s" w:sz="%s" w:space="0" w:color="%s"/>'
        '<w:left w:val="%s" w:sz="%s" w:space="0" w:color="%s"/>'
        '<w:right w:val="%s" w:sz="%s" w:space="0" w:color="%s"/>'
        '<w:bottom w:val="%s" w:sz="%s" w:space="0" w:color="%s"/>'
        '</w:tcBorders>' % (
            nsdecls('w'),
            border_style, border_size, border_color,
            border_style, border_size, border_color,
            border_style, border_size, border_color,
            border_style, border_size, border_color,
        )
    )

    THIN = "4"

    def _cell_border_xml(r_idx, c_idx, is_card_boundary):
        """Build border XML for a data cell.
        Outer edges use config table_border (style/size/color);
        interior borders use thin single black."""
        is_top    = r_idx == 0
        is_bottom = False  # footer row handles the outer bottom border
        is_left   = c_idx == 0
        is_right  = c_idx == NUM_COLUMNS - 1

        top_style    = border_style if is_top    else "single"
        top_sz       = border_size  if is_top    else THIN
        top_color    = border_color if is_top    else "000000"

        bottom_style = border_style if is_bottom else "single"
        bottom_sz    = border_size  if is_bottom else THIN
        bottom_color = border_color if is_bottom else "000000"

        left_style   = border_style if is_left   else "single"
        left_sz      = border_size  if is_left   else THIN
        left_color   = border_color if is_left   else "000000"

        right_style  = border_style if is_right  else "single"
        right_sz     = border_size  if is_right  else THIN
        right_color  = border_color if is_right  else "000000"

        return (
            '<w:tcBorders %s>'
            '<w:top w:val="%s" w:sz="%s" w:space="0" w:color="%s"/>'
            '<w:left w:val="%s" w:sz="%s" w:space="0" w:color="%s"/>'
            '<w:right w:val="%s" w:sz="%s" w:space="0" w:color="%s"/>'
            '<w:bottom w:val="%s" w:sz="%s" w:space="0" w:color="%s"/>'
            '</w:tcBorders>' % (
                nsdecls('w'),
                top_style, top_sz, top_color,
                left_style, left_sz, left_color,
                right_style, right_sz, right_color,
                bottom_style, bottom_sz, bottom_color,
            )
        )

    def build_loto_table(table_data, table_num, round_num):
        """Build a loto table (1 round + 1 title + 9 data + 1 footer), return XML element."""
        # --- Pick special cells: max 1 per row, only cells with numbers ---
        special_positions = set()  # set of (r_idx, c_idx)
        if special_count > 0:
            # Collect all filled positions grouped by row
            filled_by_row = {}
            for r in range(9):
                filled_by_row[r] = [c for c in range(NUM_COLUMNS) if table_data[r][c] is not None]
            # Pick up to special_count rows (without replacement), then 1 cell per row
            available_rows = [r for r in range(9) if filled_by_row[r]]
            chosen_rows = random.sample(available_rows, min(special_count, len(available_rows)))
            for r in chosen_rows:
                c = random.choice(filled_by_row[r])
                special_positions.add((r, c))

        # Footer is split into message (7 columns) and table id (2 columns).
        msg = random.choice(footer_messages)
        table_footer_text = f"{special_char} {msg}"
        table_id_text = f"{table_num:03d}"

        # Create table with only 9 data rows; title & footer added via raw XML
        tbl = doc.add_table(rows=9, cols=9)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        # Detach from document body — will be inserted into an outer cell
        doc.element.body.remove(tbl._tbl)

        # Table width: exact (columns × col_width)
        tbl_pr = tbl._tbl.tblPr
        col_width_twips = int(col_width.inches * 1440)
        tbl_width_twips = col_width_twips * NUM_COLUMNS
        tbl_pr.append(parse_xml(
            f'<w:tblW {nsdecls("w")} w:w="{tbl_width_twips}" w:type="dxa"/>'
        ))
        # Fixed table layout — prevents auto-resizing
        tbl_pr.append(parse_xml(
            f'<w:tblLayout {nsdecls("w")} w:type="fixed"/>'
        ))
        # Minimal cell margins
        tbl_pr.append(parse_xml(
            f'<w:tblCellMar {nsdecls("w")}>'
            '<w:top w:w="0" w:type="dxa"/>'
            '<w:left w:w="0" w:type="dxa"/>'
            '<w:bottom w:w="0" w:type="dxa"/>'
            '<w:right w:w="0" w:type="dxa"/>'
            '</w:tblCellMar>'
        ))
        # Set explicit column widths via tblGrid
        tbl_grid = tbl._tbl.find(qn('w:tblGrid'))
        if tbl_grid is None:
            tbl_grid = parse_xml(f'<w:tblGrid {nsdecls("w")}/>')
            tbl._tbl.insert(1, tbl_grid)
        else:
            for child in list(tbl_grid):
                tbl_grid.remove(child)
        for _ in range(NUM_COLUMNS):
            tbl_grid.append(parse_xml(
                f'<w:gridCol {nsdecls("w")} w:w="{col_width_twips}"/>'
            ))

        # --- Row 0: Title (built from raw XML with gridSpan, inserted before data rows) ---
        title_row_twips = int(row_ht.inches * 1440)
        header_label = title_text.upper() if header_caps else title_text
        header_font_sz_hps = int(header_font_size.pt * 2)
        title_tr = parse_xml(
            '<w:tr %s>'
            '<w:trPr>'
            '<w:trHeight w:val="%d" w:hRule="exact"/>'
            '</w:trPr>'
            '<w:tc>'
            '<w:tcPr>'
            '<w:tcW w:w="%d" w:type="dxa"/>'
            '<w:gridSpan w:val="9"/>'
            '<w:vAlign w:val="center"/>'
            '<w:shd w:fill="%s" w:val="clear" w:color="auto"/>'
            '<w:tcBorders>'
            '<w:top w:val="%s" w:sz="%s" w:space="0" w:color="%s"/>'
            '<w:left w:val="%s" w:sz="%s" w:space="0" w:color="%s"/>'
            '<w:right w:val="%s" w:sz="%s" w:space="0" w:color="%s"/>'
            '<w:bottom w:val="%s" w:sz="%s" w:space="0" w:color="%s"/>'
            '</w:tcBorders>'
            '</w:tcPr>'
            '<w:p>'
            '<w:pPr>'
            '<w:jc w:val="center"/>'
            '<w:spacing w:before="0" w:after="0" w:line="%d" w:lineRule="exact"/>'
            '</w:pPr>'
            '<w:r>'
            '<w:rPr>'
            '<w:rFonts w:ascii="%s" w:hAnsi="%s" w:cs="%s"/>'
            '<w:b/><w:bCs/>'
            '<w:sz w:val="%d"/><w:szCs w:val="%d"/>'
            '<w:color w:val="FFFFFF"/>'
            '</w:rPr>'
            '<w:t xml:space="preserve">%s</w:t>'
            '</w:r>'
            '</w:p>'
            '</w:tc>'
            '</w:tr>' % (
                nsdecls('w'),
                title_row_twips,
                tbl_width_twips,
                header_bg,
                border_style, border_size, border_color,
                border_style, border_size, border_color,
                border_style, border_size, border_color,
                border_style, border_size, border_color,
                title_row_twips,
                font_name, font_name, font_name,
                header_font_sz_hps, header_font_sz_hps,
                header_label,
            )
        )
        # Insert title row before the first data row
        first_data_tr = tbl._tbl.find(qn('w:tr'))
        first_data_idx = list(tbl._tbl).index(first_data_tr)
        tbl._tbl.insert(first_data_idx, title_tr)

        # --- Round row (inserted before the title row) ---
        round_label = f"Round {round_num}"
        round_font_sz_hps = int(header_font_size.pt * 2)
        round_tr = parse_xml(
            '<w:tr %s>'
            '<w:trPr>'
            '<w:trHeight w:val="%d" w:hRule="exact"/>'
            '</w:trPr>'
            '<w:tc>'
            '<w:tcPr>'
            '<w:tcW w:w="%d" w:type="dxa"/>'
            '<w:gridSpan w:val="9"/>'
            '<w:vAlign w:val="center"/>'
            '<w:tcBorders>'
            '<w:top w:val="%s" w:sz="%s" w:space="0" w:color="%s"/>'
            '<w:left w:val="%s" w:sz="%s" w:space="0" w:color="%s"/>'
            '<w:right w:val="%s" w:sz="%s" w:space="0" w:color="%s"/>'
            '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            '</w:tcBorders>'
            '</w:tcPr>'
            '<w:p>'
            '<w:pPr>'
            '<w:jc w:val="center"/>'
            '<w:spacing w:before="0" w:after="0" w:line="%d" w:lineRule="exact"/>'
            '</w:pPr>'
            '<w:r>'
            '<w:rPr>'
            '<w:rFonts w:ascii="%s" w:hAnsi="%s" w:cs="%s"/>'
            '<w:b/><w:bCs/>'
            '<w:sz w:val="%d"/><w:szCs w:val="%d"/>'
            '<w:color w:val="000000"/>'
            '</w:rPr>'
            '<w:t xml:space="preserve">%s</w:t>'
            '</w:r>'
            '</w:p>'
            '</w:tc>'
            '</w:tr>' % (
                nsdecls('w'),
                title_row_twips,
                tbl_width_twips,
                border_style, border_size, border_color,
                border_style, border_size, border_color,
                border_style, border_size, border_color,
                title_row_twips,
                font_name, font_name, font_name,
                round_font_sz_hps, round_font_sz_hps,
                round_label,
            )
        )
        # Insert round row at the very top (before title)
        first_tr = tbl._tbl.find(qn('w:tr'))
        tbl._tbl.insert(list(tbl._tbl).index(first_tr), round_tr)

        # --- Rows: Data (rows already created by add_table, offset by 2 for round + title) ---
        row_ht_twips = int(row_ht.inches * 1440)
        for r_idx in range(9):
            row = tbl.rows[r_idx + 2]
            row.height = row_ht
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            # Force exact row height (no auto-grow)
            tr_pr = row._tr.get_or_add_trPr()
            tr_pr.append(parse_xml(
                f'<w:trHeight {nsdecls("w")} w:val="{row_ht_twips}" w:hRule="exact"/>'
            ))
            is_card_boundary = (r_idx + 1) % 3 == 0

            for c_idx in range(NUM_COLUMNS):
                cell = row.cells[c_idx]
                cell.width = col_width
                # Force exact cell width
                tc_pr_w = cell._tc.get_or_add_tcPr()
                tc_pr_w.append(parse_xml(
                    f'<w:tcW {nsdecls("w")} w:w="{col_width_twips}" w:type="dxa"/>'
                ))
                # Vertical center
                tc_pr_w.append(parse_xml(
                    f'<w:vAlign {nsdecls("w")} w:val="center"/>'
                ))
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.space_before = Pt(0)
                p.space_after = Pt(0)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = row_ht
                p.paragraph_format.line_spacing_rule = 4  # EXACTLY

                val = table_data[r_idx][c_idx]
                is_special = (r_idx, c_idx) in special_positions
                if val is not None:
                    display_val = special_char if (is_special and special_replace) else str(val)
                    run = p.add_run(display_val)
                    run.font.size = font_size
                    run.font.name = font_name
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0, 0, 0)
                else:
                    run = p.add_run("")
                    run.font.size = font_size
                    run.font.name = font_name
                    cell._tc.get_or_add_tcPr().append(
                        parse_xml(f'<w:shd {nsdecls("w")} w:fill="F0F0F0"/>')
                    )

                tc_pr = cell._tc.get_or_add_tcPr()
                tc_pr.append(parse_xml(_cell_border_xml(r_idx, c_idx, is_card_boundary)))

        # --- Footer row (appended via raw XML with gridSpan) ---
        footer_ht_twips = int(row_ht.inches * 1440)
        footer_font_sz_hps = int(footer_font_size.pt * 2)  # half-points
        footer_tr = parse_xml(
            '<w:tr %s>'
            '  <w:trPr>'
            '    <w:trHeight w:val="%d" w:hRule="exact"/>'
            '  </w:trPr>'
            '  <w:tc>'
            '    <w:tcPr>'
            '      <w:tcW w:w="%d" w:type="dxa"/>'
            '      <w:gridSpan w:val="7"/>'
            '      <w:vAlign w:val="center"/>'
            '      <w:tcBorders>'
            '        <w:top w:val="single" w:sz="%s" w:space="0" w:color="000000"/>'
            '        <w:left w:val="%s" w:sz="%s" w:space="0" w:color="%s"/>'
            '        <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            '        <w:bottom w:val="%s" w:sz="%s" w:space="0" w:color="%s"/>'
            '      </w:tcBorders>'
            '    </w:tcPr>'
            '    <w:p>'
            '      <w:pPr>'
            '        <w:jc w:val="center"/>'
            '        <w:spacing w:before="0" w:after="0" w:line="%d" w:lineRule="exact"/>'
            '      </w:pPr>'
            '      <w:r>'
            '        <w:rPr>'
            '          <w:rFonts w:ascii="%s" w:hAnsi="%s" w:cs="%s"/>'
            '          <w:i/>'
            '          <w:iCs/>'
            '          <w:sz w:val="%d"/>'
            '          <w:szCs w:val="%d"/>'
            '          <w:color w:val="000000"/>'
            '        </w:rPr>'
            '        <w:t xml:space="preserve">%s</w:t>'
            '      </w:r>'
            '    </w:p>'
            '  </w:tc>'
            '  <w:tc>'
            '    <w:tcPr>'
            '      <w:tcW w:w="%d" w:type="dxa"/>'
            '      <w:gridSpan w:val="2"/>'
            '      <w:vAlign w:val="center"/>'
            '      <w:tcBorders>'
            '        <w:top w:val="single" w:sz="%s" w:space="0" w:color="000000"/>'
            '        <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            '        <w:right w:val="%s" w:sz="%s" w:space="0" w:color="%s"/>'
            '        <w:bottom w:val="%s" w:sz="%s" w:space="0" w:color="%s"/>'
            '      </w:tcBorders>'
            '    </w:tcPr>'
            '    <w:p>'
            '      <w:pPr>'
            '        <w:jc w:val="center"/>'
            '        <w:spacing w:before="0" w:after="0" w:line="%d" w:lineRule="exact"/>'
            '      </w:pPr>'
            '      <w:r>'
            '        <w:rPr>'
            '          <w:rFonts w:ascii="%s" w:hAnsi="%s" w:cs="%s"/>'
            '          <w:i/>'
            '          <w:iCs/>'
            '          <w:sz w:val="%d"/>'
            '          <w:szCs w:val="%d"/>'
            '          <w:color w:val="000000"/>'
            '        </w:rPr>'
            '        <w:t xml:space="preserve">%s</w:t>'
            '      </w:r>'
            '    </w:p>'
            '  </w:tc>'
            '</w:tr>' % (
                nsdecls('w'),
                footer_ht_twips,
                col_width_twips * 7,
                THIN,
                border_style, border_size, border_color,
                border_style, border_size, border_color,
                footer_ht_twips,
                font_name, font_name, font_name,
                footer_font_sz_hps,
                footer_font_sz_hps,
                table_footer_text,
                col_width_twips * 2,
                THIN,
                border_style, border_size, border_color,
                border_style, border_size, border_color,
                footer_ht_twips,
                font_name, font_name, font_name,
                footer_font_sz_hps,
                footer_font_sz_hps,
                table_id_text,
            )
        )
        tbl._tbl.append(footer_tr)

        return tbl._tbl

    # --- Build document page by page ---
    for page_start in range(0, len(all_tables), tables_per_page):
        page_tables = all_tables[page_start:page_start + tables_per_page]
        num_pairs = (len(page_tables) + COLS_PER_ROW - 1) // COLS_PER_ROW

        # Outer layout table (invisible borders, 2 columns)
        outer = doc.add_table(rows=num_pairs, cols=COLS_PER_ROW)
        outer.alignment = WD_TABLE_ALIGNMENT.CENTER
        outer_pr = outer._tbl.tblPr
        outer_pr.append(parse_xml(
            f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>'
        ))
        # Remove all outer table borders
        outer_pr.append(parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            '<w:top w:val="none" w:sz="0" w:space="0"/>'
            '<w:left w:val="none" w:sz="0" w:space="0"/>'
            '<w:right w:val="none" w:sz="0" w:space="0"/>'
            '<w:bottom w:val="none" w:sz="0" w:space="0"/>'
            '<w:insideH w:val="none" w:sz="0" w:space="0"/>'
            '<w:insideV w:val="none" w:sz="0" w:space="0"/>'
            '</w:tblBorders>'
        ))
        # Spacing between loto tables
        outer_pr.append(parse_xml(
            f'<w:tblCellMar {nsdecls("w")}>'
            '<w:top w:w="80" w:type="dxa"/>'
            '<w:left w:w="80" w:type="dxa"/>'
            '<w:bottom w:w="80" w:type="dxa"/>'
            '<w:right w:w="80" w:type="dxa"/>'
            '</w:tblCellMar>'
        ))

        # Place loto tables into the grid
        for idx, table_data in enumerate(page_tables):
            r = idx // COLS_PER_ROW
            c = idx % COLS_PER_ROW
            table_num = page_start + idx + 1

            cell = outer.cell(r, c)
            # Remove borders from outer cell
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_pr.append(parse_xml(NO_BORDER_XML))
            tc_pr.append(parse_xml(
                f'<w:vAlign {nsdecls("w")} w:val="top"/>'
            ))
            # Remove the default paragraph from the outer cell entirely
            for child_p in cell._tc.findall(qn('w:p')):
                cell._tc.remove(child_p)

            # Build and insert nested loto table
            round_num = ((table_num - 1) // tables_per_round) + 1
            if round_num > num_rounds:
                round_num = num_rounds
            inner_elem = build_loto_table(table_data, table_num, round_num)
            cell._tc.append(inner_elem)
            # Word requires at least one <w:p> in each cell; add a zero-height trailing one
            cell._tc.append(parse_xml(
                '<w:p %s>'
                '<w:pPr><w:spacing w:before="0" w:after="0" w:line="0" w:lineRule="exact"/>'
                '<w:rPr><w:sz w:val="2"/></w:rPr></w:pPr>'
                '</w:p>' % nsdecls('w')
            ))

        # Handle empty outer cells on last page
        for idx in range(len(page_tables), num_pairs * COLS_PER_ROW):
            r = idx // COLS_PER_ROW
            c = idx % COLS_PER_ROW
            cell = outer.cell(r, c)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_pr.append(parse_xml(NO_BORDER_XML))
            # Replace default paragraph with zero-height one
            for child_p in cell._tc.findall(qn('w:p')):
                cell._tc.remove(child_p)
            cell._tc.append(parse_xml(
                '<w:p %s>'
                '<w:pPr><w:spacing w:before="0" w:after="0" w:line="0" w:lineRule="exact"/>'
                '<w:rPr><w:sz w:val="2"/></w:rPr></w:pPr>'
                '</w:p>' % nsdecls('w')
            ))

    # --- Clean up: remove all auto-generated paragraphs between tables ---
    # python-docx inserts a paragraph before each table; remove them
    body = doc.element.body
    # Collect all top-level elements
    elements = list(body)
    # Remove all paragraphs (w:p) from body — they are just spacers
    for el in elements:
        if el.tag == qn('w:p'):
            body.remove(el)

    doc.save(path)


def main(num_tables, tables_per_page=8, config_path=None):
    # --- Load config ---
    if config_path is None:
        config_path = os.path.join(os.path.dirname(__file__), "config.json")
    if os.path.exists(config_path):
        with open(config_path, "r", encoding="utf-8") as f:
            config = json.load(f)
        print(f"Config loaded: {config_path}")
    else:
        config = {"title": "LÔ TÔ"}
        print(f"Config not found at {config_path}, using defaults.")

    # Initialize column ranges from config
    num_cols = config.get("columns", 9)
    init_column_ranges(num_cols)

    # Use total_tables from config if -n not explicitly provided
    if num_tables is None:
        num_tables = config.get("total_tables", 500)

    start = time.perf_counter()
    total = num_tables
    workers = NUM_WORKERS
    batch_size = (total + workers - 1) // workers  # ceil division

    print(f"Generating {total} unique tables using {workers} workers...")

    all_tables = []
    global_seen = set()

    with ProcessPoolExecutor(max_workers=workers) as executor:
        # Distribute work evenly across workers
        batches = []
        remaining = total
        for _ in range(workers):
            b = min(batch_size, remaining)
            if b > 0:
                batches.append(b)
                remaining -= b

        futures = {executor.submit(worker_batch, b): b for b in batches}

        for future in as_completed(futures):
            tables = future.result()
            for t in tables:
                fp = table_fingerprint(t)
                if fp not in global_seen:
                    global_seen.add(fp)
                    all_tables.append(t)

    # If cross-worker duplicates reduced our count, generate more
    while len(all_tables) < total:
        t = generate_table()
        fp = table_fingerprint(t)
        if fp not in global_seen:
            global_seen.add(fp)
            all_tables.append(t)

    all_tables = all_tables[:total]
    elapsed = time.perf_counter() - start
    print(f"Generated {len(all_tables)} unique tables in {elapsed:.2f}s")

    # --- Validate all ---
    invalid = 0
    for i, table in enumerate(all_tables):
        if not validate_table(table):
            invalid += 1
            print(f"  [FAIL] Table {i+1}")
    if invalid == 0:
        print(f"Validation: ALL {len(all_tables)} tables PASSED ✓")
    else:
        print(f"Validation: {invalid} tables FAILED")

    # --- Save visual output ---
    visual_path = os.path.join(os.path.dirname(__file__), "loto_tables.txt")
    with open(visual_path, "w", encoding="utf-8") as f:
        for i, table in enumerate(all_tables):
            print_combined_visual(table, i + 1, file=f)
            f.write("\n")
    print(f"Visual output → {visual_path}")

    # --- Save JSON output ---
    json_path = os.path.join(os.path.dirname(__file__), "loto_tables.json")
    json_data = []
    for i, table in enumerate(all_tables):
        json_data.append({
            "table_id": f"table-{i+1:03d}",
            "grid": table
        })
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(json_data, f, indent=2)
    print(f"JSON output  → {json_path}")

    # --- Save DOCX output ---
    timestamp = time.strftime("%Y%m%d_%H%M%S")
    docx_path = os.path.join(os.path.dirname(__file__), f"loto_tables_{timestamp}.docx")
    save_docx(all_tables, docx_path, config=config, tables_per_page=tables_per_page)
    print(f"DOCX output  → {docx_path}")

    # --- Print first 3 tables as preview ---
    print(f"\n--- Preview (first 3 of {total}) ---")
    for i in range(min(3, len(all_tables))):
        print_combined_visual(all_tables[i], i + 1)
        print()


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Vietnamese Lô Tô table generator")
    parser.add_argument("-n", "--tables", type=int, default=None,
                        help="Number of unique tables to generate (default: from config or 500)")
    parser.add_argument("-p", "--per-page", type=int, default=8,
                        help="Tables per page in DOCX output (default: 8)")
    parser.add_argument("-c", "--config", type=str, default=None,
                        help="Path to config JSON file (default: config.json)")
    args = parser.parse_args()
    main(args.tables, tables_per_page=args.per_page, config_path=args.config)
